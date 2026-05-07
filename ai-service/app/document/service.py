from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from typing import Optional
from app.document.model import HrDocument
from app.core.openai import get_embedding
from app.core.pinecone import upsert_vectors, delete_vectors
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document
import io
import os
import re
import logging

logger = logging.getLogger(__name__)


async def upload_document(
        file: UploadFile,
        company_id: str,
        db: Session,
        layer: str = "hr_uploaded") -> HrDocument:

    ALLOWED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md"]
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="지원하지 않는 파일 형식입니다. (pdf, docx, txt, md)")

    content = await file.read()

    MAX_FILE_SIZE = 10 * 1024 * 1024
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail="파일 크기는 10MB 이하여야 합니다.")

    if not file.filename or file.filename.strip() == "":
        raise HTTPException(
            status_code=400,
            detail="파일명이 없습니다.")

    text = extract_text(file.filename, content)

    if not text or text.strip() == "":
        raise HTTPException(
            status_code=400,
            detail="문서에서 텍스트를 추출할 수 없습니다.")

    document = HrDocument(
        company_id=company_id,
        document_name=file.filename,
        content=text,
        layer=layer
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    try:
        await process_document(
            document.id, company_id,
            file.filename, text, layer)
    except Exception as e:
        logger.error(f"[upload_document] 처리 실패: {e}")
        db.delete(document)
        db.commit()
        raise HTTPException(
            status_code=500,
            detail=f"문서 처리 실패: {str(e)}")

    return document


def extract_text(filename: str, content: bytes) -> str:
    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(
            page.extract_text() for page in reader.pages)
    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        return "\n".join(para.text for para in doc.paragraphs)
    elif filename.endswith(".txt") or filename.endswith(".md"):
        return content.decode("utf-8")
    else:
        raise HTTPException(
            status_code=400,
            detail="지원하지 않는 파일 형식입니다.")


def parse_document_structure(text: str) -> list[dict]:
    sections = []

    pattern = r'\[([^\]]+)\]\s+([^\n]+)(.*?)(?=\n\s*\[[^\]]+\]\s+|\Z)'

    matches = list(re.finditer(pattern, text, re.DOTALL))

    if not matches:
        logger.info(
            "[parse_document_structure] 구조 없음, 전체를 단일 섹션으로"
        )
        return [{
            "category": "일반",
            "subcategory": "문서 전체",
            "full_content": text.strip(),
            "keywords": extract_keywords(text)
        }]

    for match in matches:
        category = match.group(1).strip()
        subcategory = match.group(2).strip()
        body = match.group(3).strip()

        full_content = f"[{category}] {subcategory}\n\n{body}"

        sections.append({
            "category": category,
            "subcategory": subcategory,
            "full_content": full_content,
            "keywords": extract_keywords(full_content)
        })

    logger.info(
        f"[parse_document_structure] {len(sections)}개 섹션 추출: "
        f"{[(s['category'], s['subcategory']) for s in sections]}"
    )
    return sections


def extract_keywords(text: str) -> list[str]:
    keyword_match = re.search(
        r'■\s*관련\s*키워드\s*\n([^\n■]+)',
        text
    )
    if keyword_match:
        keywords_str = keyword_match.group(1).strip()
        keywords = [
            k.strip()
            for k in re.split(r'[,、\s]+', keywords_str)
            if k.strip() and len(k.strip()) > 1
        ]
        return keywords[:15]

    return []


def chunk_section(section: dict, max_chunk_size: int = 1500) -> list[str]:
    content = section["full_content"]

    if len(content) <= max_chunk_size:
        return [content]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=[
            "\n\n■",
            "\n\n",
            "\n",
            "。",
            ". ",
            " ",
            ""
        ],
        length_function=len,
    )
    return splitter.split_text(content)


def preprocess_text(text: str) -> str:
    result = re.sub(r'\n{3,}', '\n\n', text)
    return result


async def process_document(
        document_id: str,
        company_id: str,
        document_name: str,
        text: str,
        layer: str = "hr_uploaded"):

    logger.info(
        f"[process_document] 시작: {document_name}, "
        f"{len(text)}자, layer={layer}"
    )

    preprocessed_text = preprocess_text(text)
    sections = parse_document_structure(preprocessed_text)

    if not sections:
        raise ValueError("문서에서 섹션을 추출하지 못했습니다")

    vectors = []
    global_chunk_index = 0

    for section_idx, section in enumerate(sections):
        chunks = chunk_section(section, max_chunk_size=1500)

        for chunk_idx, chunk_text_content in enumerate(chunks):
            if len(chunk_text_content.strip()) < 20:
                continue

            embedding = get_embedding(chunk_text_content)

            vectors.append({
                "id": f"{document_id}-{global_chunk_index}",
                "values": embedding,
                "metadata": {
                    "document_id": str(document_id),
                    "company_id": str(company_id),
                    "document_name": document_name,
                    "content": chunk_text_content,
                    "category": section["category"],
                    "subcategory": section["subcategory"],
                    "keywords": section["keywords"],
                    "chunk_index": global_chunk_index,
                    "section_index": section_idx,
                    "chunk_in_section": chunk_idx,
                    "layer": layer,
                }
            })
            global_chunk_index += 1

    if not vectors:
        raise ValueError("유효한 청크가 생성되지 않았습니다")

    logger.info(
        f"[process_document] 총 {len(vectors)}개 청크 생성, Pinecone 저장 중"
    )

    upsert_vectors(vectors)

    logger.info(f"[process_document] 완료: {document_name}")


def get_documents(
        company_id: str,
        db: Session,
        layer: Optional[str] = None) -> list[HrDocument]:
    """
    문서 목록 조회.

    Args:
        layer: 'platform'/'db_sync'/'hr_uploaded'/'all'. 기본값 None → 'hr_uploaded'
    """
    query = db.query(HrDocument).filter(
        HrDocument.company_id == company_id,
        HrDocument.del_yn == "NO"
    )

    if layer == "all":
        pass
    elif layer is None:
        query = query.filter(HrDocument.layer == "hr_uploaded")
    else:
        query = query.filter(HrDocument.layer == layer)

    return query.order_by(HrDocument.created_at.desc()).all()


def delete_document(
        document_id: str,
        company_id: str,
        db: Session):

    document = db.query(HrDocument).filter(
        HrDocument.id == document_id,
        HrDocument.company_id == company_id,
        HrDocument.del_yn == "NO"
    ).first()

    if not document:
        raise HTTPException(
            status_code=404,
            detail="문서를 찾을 수 없습니다.")

    # 플랫폼 문서는 삭제 불가
    if document.layer == "platform":
        raise HTTPException(
            status_code=403,
            detail="플랫폼 문서는 삭제할 수 없습니다.")

    try:
        delete_vectors(document_id)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Pinecone 삭제 실패: {str(e)}")

    document.del_yn = "YES"
    db.commit()